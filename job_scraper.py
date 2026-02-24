import sys
import os
from dotenv import load_dotenv
import aiohttp
import json
from docx import Document
import asyncio
import re

# Modules
from pull_data import DataPuller
from app.job_search_agent import JobSearchAgent
from app.job_summary_review_agent import JobSumReviewAgent
from app.location_translater import SQL_modder
from app.call_Ollama import OllamaClient

load_dotenv()

#==================== Helper Functions ====================

# Previous Titles, skills, and resume
def load_resume_as_text():
    # doc = Document(input("Provide the path for the resume file to use: "))
    doc = Document(os.getenv("RESUME"))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    return full_text

def scrape_single_job_board(new_data, company_url, company=""):
    company_list = [] # list for all jobs in new_data
    # print(f"new data Status code: {new_data['status_code']} | company: {company}")
    # print(f"Company URL: {company_url} | New Data: {new_data}")
    for item in new_data['data']:
        if type(item) is not dict:
            continue
        if not item.get('title') or item['title'] == [None] or item['title'] == "":
            continue
        link = item.get('link')
        if link and not link.startswith(("http","https")):
            link = f"{company_url}{link}"
        maker = {
            "company": item['company'] if item.get('company') is not None else company,
            "company_url": company_url,
            "title": item['title'],
            "flexibility": item['flexibility'] if item.get('flexibility') is not None else "NA",
            "url": link,
            "source":new_data['source'] if new_data.get('source') else ""
            }
        if item.get('location') is not None:
            if re.search(r"location", item['location'], re.IGNORECASE):
                item['location'] = re.sub(r'location', "", item['location'], flags=re.IGNORECASE)
            maker["location"] = item['location']
        company_list.append(maker)
    # print(f"Company List: {company_list}")
    return company_list

def scrape_multi_job_board(new_data, company_url, company):
    full_list = [] # list for all pages combined together
    # print(f"new data: {new_data}")
    adjusted_nd = {
        "data":[],
        "status":200,
        "success":True,
        "source":new_data['source']
    }
    for item in new_data['data']:
        adjusted_nd['data'].append(item['jobs'])
    # print(f"Adjusted Data: {adjusted_nd}")
    full_list += scrape_single_job_board(adjusted_nd, company_url, company)
        # print(f"page # {page}: {new_data['pages'][page]}")
        # print(f"page: {page} | current Full List: {full_list}\n\n")
    return full_list

async def scrape_sites(i, company_url, dp):
    # print(f"Payload: {i}")
    new_data = await dp.scrape_data(i['strategy'], api_method=i['api_method'])
    new_data['source'] = i['strategy']['source']
    # print(f"New Data source: {new_data['source']}. New Data: {str(new_data)}...")
    # if new_data['data'].get(0) is None:
    #     print(f"New Data: {new_data}")
    if new_data['status_code'] != 200:
        print(f"Scraping data failed. Error code {new_data['status_code']}. Error: {new_data['error'] if new_data.get('error') is not None else 'No error message provided.'}")
        return
    if 'data' not in new_data:
        print(f"Warning: No 'data' key in response from {i['company']}. Response: {new_data}")
        return
    if len(new_data['data']) != 0:
        try:
            
                maker = scrape_multi_job_board(new_data, company_url, i['company']) if new_data['data'][0].get('jobs') is not None else scrape_single_job_board(new_data, company_url, i['company'])
        except (KeyError, IndexError, TypeError) as e:
            print(f"Error: {e}\nNew Data: {new_data}")
    else:
        return
    # print(f"Maker: {maker}")
    return maker



#==================== AI Functions ========================

# Call on Ollama. Send with system message to rate the job title and return that back with a score in JSON format
async def job_title_score(client, jt, pt, sk):
    job_title = jt
    try:
        ts = await client.title_review(job_title, pt, sk)
        return ts
    except Exception as e:
        print("❌ Error calling Ollama:", e)

# run through getting the unranked titles and scoring them
async def rate_titles(sql_queries, prev_titles, skills, dp):
        # Get unranked titles
    titles_to_process = dp.pull_data_db(sql_queries['unranked_titles'])
    print(f"Total Titles to process: {len(titles_to_process)}")

    # Run AI over unranked titles, or just score them
    score_list = []
    agent = await JobSearchAgent.create()
    for i in titles_to_process:
        # print(f"i: {i}")
        score_list.append({
            'id': i['id'],
            'score': await job_title_score(agent, i['job_name'], prev_titles, skills)
        })

    # Update the titles with the title rankings
    for i in score_list:
        query = f'''
        UPDATE job
        SET title_rating = {i['score']}
        WHERE id = {i['id']}
        '''
        dp.insert_data_db(query)
        print(f"Updated job id {i['id']} with score {i['score']}")
    dp.commit_data_db()
    print("Title scores updated in DB.")

    # unload Ollama model and close connection
    await agent.clear_ollama()
    await agent.close()

    return

async def get_job_payrate(client, job_description): 
    # print(f"sum data sum: {summary['data']['summary']}")
    pay_range = await get_pay_range(client, job_description)
    print(f"AI response for pay: {pay_range}")
    return pay_range


# Call on Ollama. Send with system message to rate the job title and return that back with a score in JSON format
async def job_summary_score(client, js, res):
    job_summary = js
    # print(f"Job summary: {str(job_summary)[:30]}... | Resume: {str(res)[:30]}... | Client: {client}")
    try:
        jsum = await client.summary_review(job_summary, res)
        return jsum
    except Exception as e:
        print("❌ Error calling Ollama:", e)
        
async def get_pay_range(client, js):
    schema = {
        "type": "object",
        "properties": {
            "pay_range": {"type": "string"}
        },
        "required": ["pay_range"]
    }
    response = await client.call(str(js), sch=schema)
    print(f"Raw response for pay range extraction: {response['pay_range']}")
    return response['pay_range']

async def fix_payrange_main():
    scrape_summaries = f'''
        SELECT j.id,j.link, c.company_name AS company, j.source AS source, j.job_summary AS summary, j.pay_range as pay_range
        FROM job j
        JOIN company c ON j.company_id = c.id
        WHERE j.title_rating >= 80 AND j.skip IS NOT True AND j.pay_range IS NULL;
    '''
    # Load .env variables
    # prev_titles = json.loads(os.getenv("PREV_TITLES", "[]"))
    # skills = json.loads(os.getenv("SKILLS", "[]"))

    # Create Data Puller Object
    dp = DataPuller(
        dbname = os.getenv("DB_NAME", ""),
        user = os.getenv("DB_USER", ""),
        password = os.getenv("DB_PASSWORD", ""),
        host = os.getenv("DB_HOST", "localhost"),
        port = os.getenv("DB_PORT", "5432")
    )
    pull_ss_list = dp.pull_data_db(scrape_summaries) # pull in the jobs that need summaries scraped
    scrape_sum_list = [dict(row) for row in pull_ss_list]  #presuming data in list

    data = []

    for i in scrape_sum_list:
        job_description = {
            'id': i['id'],
            'summary': i['summary'],
            'pay': i['pay_range'] if i.get('pay_range') is not None else None
        }
        if job_description.get('pay') is None:
            session = aiohttp.ClientSession()
            sys_message = """
            You will be provided with a job description. Your task is to extract the pay range mentioned in the description. If a pay range is found, return it in the 
            format 'X - Y' where X is the minimum pay and Y is the maximum pay. If multiple pay ranges are mentioned, return a new, single range with the smallest pay 
            to the largest pay. Ex:
            Input:
            Zone A: $135k-$174k
            Zone B: $121k-$163k
            Zone C: $115k-$152k
            Germany: €98k-€141k

            Output:
            $115k-$174k

            If no pay range is mentioned, respond with 'Not specified'. Only return pay ranges in USD as shown in the job description. Ensure 
            that your response is concise and only contains the pay range or 'Not specified'."""
            client = OllamaClient(session=session, model="hf.co/bartowski/nvidia_Orchestrator-8B-GGUF:Q4_K_M", system_message=sys_message)
            job_description['pay'] = await get_job_payrate(client, i['summary'])
            await client.unload()
            await client.session.close()

        data.append(job_description)
    print(f"Total Summaries scraped: {len(data)}")

    # Update the scraped summaries in the DB
    for i in data:
        summary = i['summary']
        id = i['id']
        if i.get('pay') != 'Not specified' and (i.get('pay') is not None or i.get('pay') != None):
            pay = i['pay']
            query = f'''
            UPDATE job
            SET pay_range = %s
            WHERE id = %s
            '''
        else:
            continue
        dp.insert_data_db(query, (pay, id))
    dp.commit_data_db()

#==================== Job Scraper =========================

async def main():
    # Load .env variables
    prev_titles = json.loads(os.getenv("PREV_TITLES", "[]"))
    skills = json.loads(os.getenv("SKILLS", "[]"))
    resume = load_resume_as_text()

    # Create Data Puller Object
    dp = DataPuller(
        dbname = os.getenv("DB_NAME", ""),
        user = os.getenv("DB_USER", ""),
        password = os.getenv("DB_PASSWORD", ""),
        host = os.getenv("DB_HOST", "localhost"),
        port = os.getenv("DB_PORT", "5432")
    )
    
    # Get sites file from .env and pulls the sites in. Needs to be a csv set up with name and site columns
    """
    JOB_SITES should point to a csv file formatted as:
    name,site
    """
    
    sites_file = os.getenv("JOB_SITES","")
    sites = dp.load_sites_list(sites_file)
    print("Sites Retrieved")
    
    # Pull in the site strategies based on the sites pulled from the sites file
    site_strategies = []
    data = []
    for i in range(len(sites['name'])): # type: ignore
        strategy = {
            "company": sites['name'][i], # type: ignore
            "site": sites['site'][i], # type: ignore
            "strategy": dp.load_site_strategies(f"./site_strategies/{sites['name'][i]}.json"), # type: ignore
            "api_method": "",
        }

        if strategy['strategy'].get('pagination') is not None:
            strategy['api_method'] = "extract-paginated"
        elif strategy['strategy'].get("js_config") is not None:
            strategy['api_method'] = "extract-js"
        else:
            strategy['api_method'] = "extract"
        print(f"Company: {strategy['company']} | API method: {strategy['api_method']}")
        site_strategies.append(strategy)
    print("Site strategies loaded.")

    # Scrape the jobs from the sites using the link to the sites and the attached strategy. Fields should be set to return matching amount of records.
    for i in site_strategies:
        company_url = i['strategy'].pop('company_url', None)
        # print(f"i: {i}")
        print(f"Scraping {i['company']}")
        if isinstance(i['strategy']['url'],str):
            d = await scrape_sites(i, company_url, dp)
            if not d:
                continue
            if isinstance(d, list):
                data += d
            else:
                data.append(d)
        elif isinstance(i['strategy']['url'],list):
            for j in i['strategy']['url']:
                new_payload = i
                new_payload['strategy']['url'] = j
                d = await scrape_sites(new_payload, company_url, dp)
                if not d:
                    continue
                if isinstance(d, list):
                    data += d
                else:
                    data.append(d)
    print(f"Total jobs scraped: {len(data)}")
    # print(f"Data: {data}")
    # sys.exit("Stopping program")
    
    # Load in the latest pulled in jobs
    dp.load_scraped_data_to_db(data)
    
    # clear variables from data

    del data, site_strategies, sites, sites_file, strategy

    # build SQL queries
    
    sql_read_queries = {
        'unranked_titles' : f'''
        SELECT id, job_name
        FROM job
        WHERE title_rating IS NULL AND skip IS NOT True;
        ''',
        'scrape_summaries': f'''
        SELECT j.id,j.link, c.company_name AS company, j.source AS source
        FROM job j
        JOIN company c ON j.company_id = c.id
        WHERE j.title_rating >= 75 AND j.skip IS NOT True AND j.job_summary IS NULL;
        ''',
        'scrape_sum_companies': f'''
        SELECT DISTINCT c.company_name AS company, c.company_url, j.source
        FROM company c
        JOIN job j ON j.company_id = c.id
        WHERE j.title_rating >= 75 AND j.skip IS NOT True AND j.job_summary IS NULL;
        ''',
        'unranked_summaries': f'''
        SELECT j.id, j.job_summary
        FROM job j
        WHERE j.title_rating >= 75 AND skip IS NOT True AND j.job_summary IS NOT NULL and j.summary_rating IS NULL;
        '''
    }
    
    await rate_titles(sql_read_queries, prev_titles, skills, dp)
    
    # Scrape the summaries for the jobs above 80 score and update them in the DB

    pull_ss_list = dp.pull_data_db(sql_read_queries['scrape_summaries']) # pull in the jobs that need summaries scraped
    scrape_sum_list = [dict(row) for row in pull_ss_list]  #presuming data in list
    pull_companies = dp.pull_data_db(sql_read_queries['scrape_sum_companies'])
    print(f"Total Summaries to scrape: {len(scrape_sum_list)}")

    # Pull in the site strategies based on the sites pulled from the sites file
    site_strategies = {}
    data = []
    for i in pull_companies:
        print(f"Pull companies item: {i}")
        strategy = {
            "company": i[0],
            "site": i[1],
            "strategy": dp.load_site_strategies(f"./job_page_strategy/{i[2]}.json"),
            "api_method": ""
        }
        if "js_config" in strategy['strategy']:
            strategy['api_method'] = "extract-js"
        else:
            strategy['api_method'] = "extract"
        site_strategies[i[2]] = strategy
    print("Site strategies for singlejob listings loaded.")

    # Scrape the job summaries from the job posting links

    for i in scrape_sum_list:
        print(f"Company: {i['company']} | Link: {i['link']} | Source: {i['source']}")
        site_strategies[i['source']]['strategy']['url'] = i['link']
        summary = await dp.scrape_data(site_strategies[i['source']]['strategy'], api_method = site_strategies[i['source']]['api_method'])
        # print(f"Summary data: {summary}")
        if  summary.get('data') is None  or (isinstance(summary['data'], list) and len(summary['data']) > 0 and summary['data'][0].get('summary') is None):
            print(f"No summary found for job id {i['id']} at {i['link']}")
            continue
        count = 0
        # print(f"summary: {summary} | link: {i['link']} |Company: {i['company']}")
        while summary['status_code'] != 200 and count <=4:
            count += 1
            summary = await dp.scrape_data(site_strategies[i['source']]['strategy'], api_method = site_strategies[i['source']]['api_method'])
        if count >= 4:
            continue
        job_description = {
            'id': i['id'],
            'summary': summary['data'][0]['summary'],
            'pay': summary['data'][0].get('pay')
        }

        if job_description.get('pay') is None:
            session = aiohttp.ClientSession()
            sys_message = """
            You will be provided with a job description. Your task is to extract the pay range mentioned in the description. If a pay range is found, return it in the 
            format 'X - Y' where X is the minimum pay and Y is the maximum pay. If multiple pay ranges are mentioned, return a new, single range with the smallest pay 
            to the largest pay. Ex:
            Input:
            Zone A: $135k-$174k
            Zone B: $121k-$163k
            Zone C: $115k-$152k
            Germany: €98k-€141k

            Output:
            $115k-$174k
            
            If no pay range is mentioned, respond with 'Not specified'. Only return pay ranges in USD as shown in the job description. Ensure 
            that your response is concise and only contains the pay range or 'Not specified'."""
            client = OllamaClient(session=session, model="hf.co/bartowski/nvidia_Orchestrator-8B-GGUF:Q4_K_M", system_message=sys_message)
            job_description['pay'] = await get_job_payrate(client, summary['data'][0]['summary'])
            await client.unload()
            await client.session.close()

        data.append(job_description)
    print(f"Total Summaries scraped: {len(data)}")

    # Update the scraped summaries in the DB
    for i in data:
        # print(f"i: {i}")
        summary = i['summary']
        id = i['id']
        # print(f"i pay: {i.get('pay')}") if i.get('pay') != 'Not specified' and (i.get('pay') is not None or i.get('pay') != None) else print("No pay info found.")
        if i.get('pay') != 'Not specified' and (i.get('pay') is not None or i.get('pay') != None):
            pay = i['pay']
            query = f'''
            UPDATE job
            SET job_summary = %s, pay_range = %s
            WHERE id = %s
            '''
        else:
            query = f'''
            UPDATE job
            SET job_summary = %s
            WHERE id = %s
            '''
        summary_preview = summary[:30] if summary and isinstance(summary, str) else str(summary)[:30] if summary else "None"
        print(f"Updating job id {id} with summary {summary_preview}...")
        # print(f"query: {query}")
        dp.insert_data_db(query, (summary, pay, id) if i.get('pay') != 'Not specified' and (i.get('pay') is not None or i.get('pay') != None) else (summary, id))
    dp.commit_data_db()

    del scrape_sum_list, site_strategies, data

    # Get the unranked summaries
    summaries_to_process = dp.pull_data_db(sql_read_queries['unranked_summaries'])

    # Check unranked summaries for ones that did not pull and mark them to skip

    for i in range(len(summaries_to_process)):
        if summaries_to_process[i]['job_summary'] == None:
            summaries_to_process.pop(i)

    print("processing summaries...")
    # Run AI over the unranked summaries
    agent = await JobSumReviewAgent.create()
    score_list = []
    
    for i in summaries_to_process:
        score_list.append({
            'id': i['id'],
            'score': await job_summary_score(agent, i['job_summary'].split('.'), resume)
        })
        print(f"Processed summary for job id {i['id']}. Score: {str(score_list[-1]['score'])[:30]}")

    # Update the scores for the unranked summaries

    for i in score_list:
        rating = i['score'][0]
        jsr_reasoning = i['score'][1]
        id = i['id']
        query = f'''
        UPDATE job
        SET summary_rating = %s, jsr_reasoning = %s
        WHERE id = %s
        '''
        dp.insert_data_db(query, (rating, jsr_reasoning, id))
    dp.commit_data_db()
    print(f"Summary scores updated in DB. Total updated: {len(score_list)}")
    print(f"Job complete.")
    
    # unload Ollama model and close connection
    await agent.clear_ollama()
    await agent.close()

    database = SQL_modder(
        dbname = os.getenv("DB_NAME", ""),
        user = os.getenv("DB_USER", ""),
        password = os.getenv("DB_PASSWORD", ""),
        host = os.getenv("DB_HOST", "localhost"),
        port = os.getenv("DB_PORT", "5432")
    )
    empty = database.get_empty_locations()
    print(f"number of offices with empty locations: {len(empty)}")
    breakdown_locations = await database.breakdown_location(empty)
    print(f"locations broken down: {len(breakdown_locations)}")
    database.update_db(breakdown_locations)
    print(f"all done updating city/state/country")

    return

if __name__ == "__main__":
    asyncio.run(main())

