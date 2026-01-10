import json
import aiohttp
from app.call_Ollama import OllamaClient
import asyncio
from dotenv import load_dotenv
import os
from docx import Document


class JobSumReviewAgent:
    def __init__(self, client: OllamaClient):
        self.client = client

    @classmethod
    async def create(cls):
        """Async factory method for safe init"""
        SYSTEM_MESSAGE = """
            You are an expert career coach reviewing the user's resume against the description of a job to determine
            if the user would be a good fit for the job. 
    
            <INPUT>
            
            Job Description: <string>,
                Resume: <string>,
                Criterion: <string>,
                Score Range: <string> 
            
            <OUTPUT>
            
            {
                "type": "object",
                "properties": {
                    "criterion": {"type": "string"},
                    "reasoning": {"type": "string"},
                    "score": {"type": "number"}
                }
                "required": ["criterion", "reasoning", "score"]
            }
            
            For each request:
            - You will receive ONE Job description, a resume, a Criterion description, and a score range.        
            - Compare the Job description and the resume
            - Demonstrate careful analysis and consideration of multiple angles.
            - Use the given Criterion description for details on how to score the match.
            - Score must be within the given *Score Range*.
            - OUTPUT MUST BE VALID JSON, MATCHING THE <OUTPUT> SCHEMA EXACTLY.
            """
        session = aiohttp.ClientSession()
        client = OllamaClient(session, model="hermes3:8b", system_message=SYSTEM_MESSAGE) # Note: think nvidia Orchestrator and Microsoft Fara may be too kind in its ratings. Just does not like Qwen
        return cls(client)

    async def summary_review (self, js: str, res: str):
        """
        Runs a series of Ollama calls to get a score for the job description

        Args:
        js: Job description, in the form of a string
        res: User's resume, in the form of a string

        Returns:
            int: score for the Job Title
        """
        # ---------------------------
        # Prompt pieces
        # ---------------------------

        criteria = [
            {
                "message": """
                Compare the following resume and job description. Give a rating for how well the resume matches the job description out of 100%. 
                For areas where the resume does not match, give reasoning, suggestions on how to fit better, and examples for those suggestions. 
                
                Steps:
                1. Extract the required and nice to have technical skills from the job description (Qualifications/Requirements).
                2. For each skill, check if the resume has an exact match, a clear synonym/abbreviation, or same tech family. 
                Mark as MATCH or NO MATCH.
                   - Examples: AWS ↔ Amazon Web Services; JS ↔ JavaScript; Docker ↔ Containerization; REST APIs ↔ HTTP API dev.
                   - Do not match unrelated items (e.g., AWS ≠ Internal Network administration).
                   - If a skill is abstract or a soft skill, err on the user's favor for that skill
                3. Overlap % = (MATCH ÷ total skills) × 100, rounded.
                4. Score:
                   - 60–100% = 40-70
                   - 40–59% = 25-40
                   - 20–39% = 15-25
                   - 0–19% = 0-15
                Final Score: NN
                ***NOTE***  Do not worry about matching the exact verbage, or being overskilled. This resume is an overall comprehensive list of 
                skills, work, and projects. Match whether the resume shows the user has the correct experience for the role. A tailored resume can come later.
                Respond only in the specified JSON format.
                """,
                "score_range": "0-70"},
            {
                "message": "*Domain/Industry Relevance*\nIs the job in the same industry/domain as the user’s background?\nExact or related  domain (e.g. Healthcare → Biotech) → 10-15 pts\nDifferent but tangential domain (e.g. Marketing → EdTech) → 5–10 pts\nCompletely unrelated → 0–5 pts. respond only in the specified JSON format.",
                "score_range": "0-15"},
            {
                "message": "*Seniority/Level Fit* \nIs the level of the role aligned with user’s career stage?\nNear perfect fit (e.g. Mid-level → Mid-level, or Senior → Senior) → 7-10 pts\nSlightly above/below (e.g. Senior → Lead, or Senior → Mid) → 4-6 pts\nSignificantly off (e.g. Intern → Director) → 0–3 pts. respond only in the specified JSON format.",
                "score_range": "0-10"},
            {
                "message": "*Transferability Potential* \nEven if the titles/skills don’t perfectly align, could the role leverage transferable skills?\nHigh transferability → 4–5 pts\nModerate → 2–3 pts\nLow → 0–1 pts. respond only in the specified JSON format.",
                "score_range": "0-5"}
        ]

        results = []
        json_schema = {
            "type": "object",
            "properties": {
                "criterion": {"type": "string"},
                "reasoning": {"type": "string"},
                "score": {"type": "number"}
            },
            "required": ["criterion", "reasoning", "score"]
        }

        for i in criteria: # builds the prompt while iterating over the criteria, then makes the call to Ollama
            valid = False
            user_prompt = f"""
            User Input:
            Criterion: { i['message']}
            Job Description: { js }
            Resume: { res }
            Score Range: { i['score_range'] }
            Return JSON only.
            """
            invalid = False
            invalid_count = 0
            while not valid: # Verifies the reply contained good data. If not tries again
                if not invalid:
                    result = await self.client.call(user_prompt, json_schema, 0.2, 0.5, 5, mt=1024)
                else:
                    print("Invalid response, retrying...")
                    result = await self.client.call(f"The previous response did not match the outlined criteria for JSON formatting. Retry the response following the JSON schema outlined in the system message." + user_prompt, json_schema, 0.7, 0.7, 15, mt=1024)

                try:
                    if (isinstance(result.get('score'), int) or isinstance(result.get('final_score'), int)) and result.get('reasoning'):
                        print("Valid format.")
                        valid = True
                        results.append(result)
                    else:
                        print("Invalid format." + str(result))
                        invalid = True
                        invalid_count += 1
                        if invalid_count > 3:
                            print("Too many invalid responses, ending review and continuing on to next job.")
                            result = {"reasoning": "Invalid format.", "score": 0}
                            results.append(result)
                            valid = True
                            invalid_count = 0
                except KeyError as e:
                    print(e)
                    print("KeyError, retrying...")
                    continue
        total_score = 0
        reasoning = ''
        # print(f"Results: {results}")
        for j in results:
            if 'raw_text' in j: # in case the reply comes back as raw text, try to extract and convert to JSON
                print('Raw Text!')
                print(f"Here's the Raw Text: {j['raw_text']}")
                print(json.loads(f"More Raw Text: {j['raw_text']}"))
                j_converted = json.loads(j['raw_text'])
                total_score += int(j_converted['score'])
                reasoning += j_converted['reasoning'] + "\n\n"
            else:
                # print(f"Result: {str(j)[:30]}...")
                total_score += int(j.get('score')) if 'score' in j else int(j.get('final_score'))# Adds the score to the total
                reasoning += json.dumps(j['reasoning']) + "\n\n"
        print(total_score)
        return [total_score, reasoning]

    async def clear_ollama(self): # unloads the model from Ollama
        await self.client.unload()

    async def close(self):
        """Close aiohttp session when done"""
        await self.client.session.close()

async def main():
    load_dotenv()
    agent = await JobSumReviewAgent.create()
    # declare variables from .env
    # Pull the test resume file and parse it
    doc = Document(os.getenv("RESUME"))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    resume = "\n".join([line for line in full_text if line])
    # Pull the test job description file and parse it
    doc = Document(os.getenv("JOB_DESCRIPTION"))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    job_description = "\n".join([line for line in full_text if line])
    # starts reviewing process
    score_list = []
    count = 0
    while count < 10:
        count += 1
        try:
            ts = await agent.summary_review(job_description, resume)
            print(ts)
            score_list.append(ts[0])
        except Exception as e:
            print("❌ Error calling Ollama:", e)
    average = sum(score_list) / len(score_list)
    print("Average score:", average)
    await agent.clear_ollama()
    await agent.close()
    return None

if __name__ == "__main__":
    asyncio.run(main())